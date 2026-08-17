"""Generación del informe BIEC .docx sobre la plantilla corporativa de VECTUS.

Edita la plantilla existente con python-docx (no la recrea), conservando el
branding, la portada y los estilos. Se dispara SOLO a pedido (botón/endpoint),
nunca automático. Respeta la regla de oro (B.12): no eleva severidad ni inventa;
la tabla de vulnerabilidades incluye solo hallazgos confirmados con severidad
crítica/alta/media/baja. La buena postura y las áreas a validar no son
vulnerabilidades y no entran en esa tabla.
"""
import copy
import io
import os

from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "report_template", "plantilla.docx"
)

SEV_COL = {"critica": 2, "alta": 3, "media": 4, "baja": 5}  # columnas de la tabla resumen
SEV_LABEL = {"critica": "Crítica", "alta": "Alta", "media": "Media", "baja": "Baja"}
SEV_ORDER = {"critica": 0, "alta": 1, "media": 2, "baja": 3}
_VULN_SEVS = ("critica", "alta", "media", "baja")


# ─── helpers de edición ──────────────────────────────────────────────

def _merge_replace(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    full = full.replace(old, new)
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(full)
    return True


def _xml_replace(p, old, new):
    """Reemplaza texto recorriendo todos los nodos w:t (incluye hipervínculos)."""
    for t in p._p.iter():
        if t.tag.endswith("}t") and t.text and old in t.text:
            t.text = t.text.replace(old, new)


def _set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _set_cell(cell, text):
    _set_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def _block_elements(doc, heading_text, stop_texts):
    body = doc.element.body
    children = list(body)
    start = None
    for i, el in enumerate(children):
        if el.tag.endswith("}p"):
            txt = "".join(t.text or "" for t in el.iter() if t.tag.endswith("}t"))
            if heading_text in txt:
                start = i
                break
    if start is None:
        return []
    out = [children[start]]
    for el in children[start + 1:]:
        if el.tag.endswith("}p"):
            txt = "".join(t.text or "" for t in el.iter() if t.tag.endswith("}t"))
            if any(s in txt for s in stop_texts):
                break
        out.append(el)
    return out


def _fill_finding_block(elements, idx, f):
    for el in elements:
        if el.tag.endswith("}p"):
            p = Paragraph(el, None)
            t = p.text
            if "Vulnerabilidad encontrada" in t:
                _set_text(p, f"#{idx}. {f['titulo']}")
            elif t.strip() == "Descripción detallada de la vulnerabilidad encontrada":
                _set_text(p, f.get("descripcion") or "—")
            elif "AGREGAR IMAGEN" in t:
                _set_text(p, "")
            elif "URLCOMPLETA" in t:
                _xml_replace(p, "https://URLCOMPLETA", f.get("sistema") or "—")
            elif t.strip() == "CVE de referencia a las vulnerabilidades":
                _set_text(p, f.get("cve_line") or f.get("cve") or "No aplica")
            elif t.strip() == "Plan de remediación a realizar":
                _set_text(p, f.get("recomendacion") or "—")
            elif "Mas links para que el cliente" in t:
                _set_text(p, f.get("mas_info") or "—")
        elif el.tag.endswith("}tbl"):
            tb = Table(el, None)
            labelval = {
                "vector de ataque": f.get("vector", ""),
                "complejidad": f.get("complejidad", ""),
                "privilegios": f.get("privilegios", ""),
                "interacción": f.get("interaccion", ""),
                "alcance": f.get("alcance", ""),
                "confidencialidad": f.get("imp_c", ""),
                "integridad": f.get("imp_i", ""),
                "disponibilidad": f.get("imp_d", ""),
            }
            for row in tb.rows:
                cells = row.cells
                c0 = cells[0].text.replace("\u200b", "").strip()
                if c0.startswith("Severidad"):
                    _set_cell(cells[0], f"Severidad: {SEV_LABEL.get(f['severidad'], f['severidad'])}")
                elif c0.startswith("CVSS"):
                    _set_cell(cells[0], f"CVSS: {f.get('cvss') or 'N/D'}")
                elif c0.startswith("Ocurrencias"):
                    _set_cell(cells[0], f"Ocurrencias: {f.get('ocurrencias', 1)}")
                for li, vi in ((1, 2), (3, 4)):
                    if len(cells) > vi:
                        lbl = cells[li].text.replace("\u200b", "").strip().lower()
                        for key, val in labelval.items():
                            if key in lbl and val:
                                _set_cell(cells[vi], val)
                                break


# ─── construcción del contexto desde los hallazgos ───────────────────

def _host(target: str) -> str:
    t = (target or "").replace("https://", "").replace("http://", "")
    return t.split("/")[0]


def build_context(target, cliente, ip, findings):
    """findings: lista de objetos con atributos tipo Finding."""
    from app import report_catalog

    vulns = []
    for f in findings:
        if f.estado == "confirmado" and f.severidad in _VULN_SEVS:
            v = {
                "titulo": f.titulo,
                "severidad": f.severidad,
                "cvss": (str(f.cvss) if getattr(f, "cvss", None) else None),
                "ocurrencias": getattr(f, "ocurrencias", 1) or 1,
                "sistema": f.sistema_afectado or "—",
                "evidencia": f.evidencia or "",
                "cve": f.cve or "No aplica",
                "cwe": getattr(f, "cwe", None),
                "recomendacion": f.recomendacion or "—",
                "mas_info": f.mas_info or "—",
                "dedup_key": getattr(f, "dedup_key", ""),
            }
            vulns.append(report_catalog.enrich(v))
    vulns.sort(key=lambda v: SEV_ORDER.get(v["severidad"], 9))
    alcance = target + (f" (IP {ip})" if ip else "")
    return {"cliente": cliente or _host(target), "alcance": alcance, "vulnerabilidades": vulns}


def _extract_ip(findings) -> str | None:
    """Toma la IP del hallazgo de contexto CDN/origen si existe."""
    for f in findings:
        if getattr(f, "dedup_key", "") == "contexto:cdn-origen":
            ev = f.evidencia or ""
            # evidencia tipo "IP 67.227.192.19 · PTR ..."
            for tok in ev.replace("·", " ").split():
                if tok.count(".") == 3 and tok.replace(".", "").isdigit():
                    return tok
    return None


# ─── generación ──────────────────────────────────────────────────────

def generate_bytes(context) -> bytes:
    doc = Document(TEMPLATE_PATH)

    for p in doc.paragraphs:
        _merge_replace(p, "__________", context["cliente"])
        _merge_replace(p, "Gobierno de San Luis", context["cliente"])
        if p.text.strip() == "https://Pagina web a analizar":
            _set_text(p, context["alcance"])

    vulns = context["vulnerabilidades"]

    # tabla resumen (tabla 0)
    tabla = doc.tables[0]
    total_row = tabla.rows[-1]
    sample_rows = tabla.rows[1:-1]
    plantilla_tr = copy.deepcopy(sample_rows[0]._tr)
    for r in sample_rows:
        r._tr.getparent().remove(r._tr)
    tot = {"critica": 0, "alta": 0, "media": 0, "baja": 0}
    for i, f in enumerate(vulns, 1):
        tr = copy.deepcopy(plantilla_tr)
        total_row._tr.addprevious(tr)
        rr = _Row(tr, tabla)
        tot[f["severidad"]] += 1
        cells = rr.cells
        _set_text(cells[0].paragraphs[0], f"#{i}")
        _set_text(cells[1].paragraphs[0], f["titulo"])
        for sev, col in SEV_COL.items():
            _set_text(cells[col].paragraphs[0], "1" if f["severidad"] == sev else "")
        _set_text(cells[6].paragraphs[0], "1")
    tcells = total_row.cells
    for sev, col in SEV_COL.items():
        _set_text(tcells[col].paragraphs[0], str(tot[sev]))
    _set_text(tcells[6].paragraphs[0], str(len(vulns)))

    # bloques por hallazgo: clonar el primer bloque de muestra
    blockA = _block_elements(doc, "#1. Vulnerabilidad encontrada", ["Conclusión"])
    trimmed = [blockA[0]]
    for el in blockA[1:]:
        txt = "".join(t.text or "" for t in el.iter() if t.tag.endswith("}t"))
        if "Vulnerabilidad encontrada" in txt:
            break
        trimmed.append(el)
    blockA = trimmed
    anchor = blockA[0]
    for i, f in enumerate(vulns, 1):
        clones = [copy.deepcopy(el) for el in blockA]
        _fill_finding_block(clones, i, f)
        for el in clones:
            anchor.addprevious(el)
    # borrar los dos bloques de muestra originales
    for el in _block_elements(doc, "#1. Vulnerabilidad encontrada", ["Conclusión"]):
        el.getparent().remove(el)

    # conclusión
    n = len(vulns)
    principal = vulns[0]["titulo"] if vulns else "sin vulnerabilidades confirmadas"
    riesgo_txt = (
        f"El riesgo más relevante encontrado corresponde a {principal}, que representa "
        "una amenaza directa para la seguridad de los activos evaluados (p. ej. acceso no "
        "autorizado, pérdida de datos sensibles o indisponibilidad de servicio)."
        if vulns else
        "No se identificaron vulnerabilidades confirmadas en este barrido. Se registran, de "
        "corresponder, observaciones de contexto y áreas a validar en profundidad."
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("El riesgo más relevante encontrado corresponde a"):
            _set_text(p, riesgo_txt)
    for p in doc.paragraphs:
        _merge_replace(p, "[N]\nvulnerabilidades", f"{n} vulnerabilidades")
        _merge_replace(p, "[N] vulnerabilidades", f"{n} vulnerabilidades")
        _merge_replace(p, "[dominios /\naplicaciones evaluadas]", context["alcance"])
        _merge_replace(p, "[dominios / aplicaciones evaluadas]", context["alcance"])
        _merge_replace(p, "[N] críticas", f"{tot['critica']} críticas")
        _merge_replace(p, "[N] altas", f"{tot['alta']} altas")
        _merge_replace(p, "[N] medias", f"{tot['media']} medias")
        _merge_replace(p, "[N] bajas", f"{tot['baja']} bajas")
        _merge_replace(p, "[vulnerabilidad\nprincipal]", principal)
        _merge_replace(p, "[vulnerabilidad principal]", principal)

    # Limpieza de la narrativa de la conclusión: quitar placeholders y derivar
    # las prioridades de los hallazgos reales (en vez del ejemplo de la plantilla).
    crit = [v for v in vulns if v["severidad"] == "critica"]
    alto = [v for v in vulns if v["severidad"] == "alta"]
    medio = [v for v in vulns if v["severidad"] in ("media", "baja")]

    def _band(items, empty):
        if not items:
            return empty
        return "Remediar: " + "; ".join(v["titulo"] for v in items[:6]) + "."

    prio = {
        "Actualizar Lodash": _band(crit, "No se identificaron vulnerabilidades críticas en este barrido."),
        "Definir CSP estricta": _band(alto, "No se identificaron vulnerabilidades altas en este barrido."),
        "Eliminar parámetros sensibles": _band(medio, "No se identificaron vulnerabilidades medias ni bajas en este barrido."),
    }
    quitar = ("Quitar unsafe-inline", "Implementar tokens anti-CSRF", "Higiene de publicación")
    impacto = {
        "[Impacto 1": "Pérdida de información, interrupción de servicio o daño reputacional.",
        "[Impacto 2": "Sanciones regulatorias, incumplimiento normativo o pérdida financiera.",
        "[Impacto 3": "Afectación a clientes internos o externos y a la confianza de socios.",
    }
    accion = {
        "[Acción 1": "Validar la aplicación de los parches y configuraciones recomendados.",
        "[Acción 2": "Realizar una nueva evaluación de verificación tras la remediación.",
        "[Acción 3": "Establecer una frecuencia de pruebas futuras (trimestral o semestral).",
    }
    a_borrar = []
    for p in doc.paragraphs:
        t = p.text.strip()
        _merge_replace(p, "[confidencialidad / integridad / disponibilidad]",
                       "la confidencialidad, integridad y disponibilidad")
        _merge_replace(p, "[confidencialidad /\nintegridad / disponibilidad]",
                       "la confidencialidad, integridad y disponibilidad")
        done = False
        for pref, val in prio.items():
            if t.startswith(pref):
                _set_text(p, val); done = True; break
        if done:
            continue
        if any(t.startswith(q) for q in quitar):
            a_borrar.append(p._p); continue
        for pref, val in impacto.items():
            if t.startswith(pref):
                _set_text(p, val); done = True; break
        if done:
            continue
        for pref, val in accion.items():
            if t.startswith(pref):
                _set_text(p, val); break
    for el in a_borrar:
        el.getparent().remove(el)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_for_scan(scan, findings) -> bytes:
    """Genera el informe para un scan a partir de sus hallazgos."""
    cliente = getattr(scan, "cliente", None)
    project = getattr(scan, "project", None)
    if not cliente and project is not None:
        cliente = getattr(project, "client", None) or getattr(project, "name", None)
    ip = _extract_ip(findings)
    ctx = build_context(scan.target, cliente, ip, findings)
    return generate_bytes(ctx)
